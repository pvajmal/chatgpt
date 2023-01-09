import pandas as pd
from collections import defaultdict
import docx
data = pd.read_excel('Data.xlsx')


add_data = defaultdict(list)

add_data['Name'].append('Muhammed Ajmal P V')
add_data['Address'].append('Silsila Mahal, Ice plant road, Koyilandy')
add_data['Phone'].append('+919048547406')
add_data['Summary'].append('''I am an IT engineer with 2 years of experience in the retail industry and a strong drive to succeed. I enjoy tackling complex challenges and am always
open to learning new technologies and tools in order to develop innovative solutions. My passion for technology and problem‑solving, combined
with my experience in the retail industry, make me a valuable addition to any team''')

data = pd.DataFrame.from_dict(add_data)
data.to_excel('DataE.xlsx', index = False)

doc = docx.Document()
para = doc.add_paragraph().add_run(data['Name'][0]).bold=True
doc.save('Resume_'+'.docx')